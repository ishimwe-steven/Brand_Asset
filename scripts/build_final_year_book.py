from pathlib import Path
import re
from copy import deepcopy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

root = Path(__file__).resolve().parents[1]
source = root / "Mupenzi-Research-Proposal_final.docx"
output = root / "Mupenzi-Final-Year-Project-Book.docx"
ps_source = (root / "scripts" / "build-final-year-book.ps1").read_text(encoding="utf-8")
match = re.search(r"\$chapterText = @'\n(.*?)\n'@", ps_source, re.S)
if not match:
    raise RuntimeError("Final-book chapter content was not found")
chapter_text = match.group(1)

doc = Document(source)

replacements = {
    "A research proposal submitted": "A final year project book submitted",
    "This Research Proposal": "This Final Year Project Book",
    "this research proposal": "this final year project book",
    "research proposal": "final year project book",
    "the proposed system": "the developed system",
    "The proposed system": "The developed system",
    "this study proposes the design and implementation": "this study designed and implemented",
    "will be conducted": "was conducted",
    "will be utilized": "was utilized",
    "will be used": "was used",
    "will involve": "involved",
    "will focus": "focused",
    "will contribute": "contributes",
    "is expected to help": "helps",
}
for paragraph in doc.paragraphs:
    updated = paragraph.text
    for old, new in replacements.items():
        updated = updated.replace(old, new)
    if updated != paragraph.text:
        paragraph.text = updated

abstract_text = (
    "This final year project designed and implemented VerifyAI, an Automated Brand Asset and Packaging "
    "Compliance System for Rwandan SME exporters. The system combines Optical Character Recognition, "
    "computer vision and a rule-based compliance engine to inspect packaging text, official logos, brand "
    "colours, barcodes, QR codes and market-specific requirements. A React client, Node.js/Express API, "
    "MySQL database and Python FastAPI analysis service support administrators, SME exporters and designers "
    "through role-controlled workflows. The implemented solution stores regulation requirements, official "
    "brand profiles, detected evidence, compliance issues, correction suggestions, export-readiness scores "
    "and PDF reports. Functional and integration testing confirmed secure account management, regulation "
    "processing, brand-scoped package verification, explainable results and report generation. VerifyAI is "
    "intended as pre-submission decision support and does not replace formal certification by competent "
    "authorities. Keywords: Brand Asset Verification, Packaging Compliance, OCR, Computer Vision, SMEs."
)
for index, paragraph in enumerate(doc.paragraphs):
    if paragraph.text.strip().upper() == "ABSTRACT":
        for candidate in doc.paragraphs[index + 1:]:
            if candidate.text.strip():
                candidate.text = abstract_text
                candidate.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                break
        break

ack = next((p for p in doc.paragraphs if "ACKNOWLEDG" in p.text.strip().upper()), None)
if ack is not None:
    heading_node = OxmlElement("w:p"); ack._p.addprevious(heading_node)
    heading = Paragraph(heading_node, ack._parent); heading.style = doc.styles["Heading 1"]
    heading.add_run("DEDICATION").bold = True; heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dedication_node = OxmlElement("w:p"); ack._p.addprevious(dedication_node)
    dedication = Paragraph(dedication_node, ack._parent)
    dedication.add_run("I dedicate this final year project to my family, whose prayers, encouragement and sacrifice sustained my education; to my lecturers and supervisor for their guidance; and to Rwandan entrepreneurs whose determination inspired the development of a practical tool for stronger brands and export-ready packaging.")
    dedication.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

toc_heading = next((p for p in doc.paragraphs if p.text.strip().upper() == "TABLE OF CONTENTS"), None)
if toc_heading is not None:
    toc_node = OxmlElement("w:p"); toc_heading._p.addnext(toc_node)
    toc_paragraph = Paragraph(toc_node, toc_heading._parent)
    run = toc_paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText"); instruction.set(qn("xml:space"), "preserve"); instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t"); placeholder.text = "Right-click and select Update Field to generate the table of contents."
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])

reference_paragraph = next((p for p in doc.paragraphs if p.text.strip().lower() in {"refersences", "references"}), None)
if reference_paragraph is None:
    raise RuntimeError("References heading was not found")

for line in chapter_text.splitlines():
    node = OxmlElement("w:p")
    reference_paragraph._p.addprevious(node)
    paragraph = Paragraph(node, reference_paragraph._parent)
    text = line.strip()
    if not text:
        continue
    run = paragraph.add_run(text)
    if text.startswith("CHAPTER "):
        paragraph.style = doc.styles["Heading 1"]
        run.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.page_break_before = True
    elif re.match(r"^\d+\.\d+\.\d+ ", text):
        paragraph.style = doc.styles["Heading 3"]
    elif re.match(r"^\d+\.\d+ ", text):
        paragraph.style = doc.styles["Heading 2"]
    elif text.startswith("Figure "):
        run.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.name = "Courier New"
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

reference_paragraph.text = "REFERENCES"
reference_paragraph.style = doc.styles["Heading 1"]
reference_paragraph.paragraph_format.page_break_before = True

for section in doc.sections:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText"); instruction.set(qn("xml:space"), "preserve"); instruction.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])

doc.core_properties.title = "Automated Brand Asset and Packaging Compliance System"
doc.core_properties.subject = "Final Year Project Book"
doc.core_properties.author = "Innocent Mupenzi"
doc.save(output)
print(output)
